import streamlit as st
import google.generativeai as genai
import user_manager
import time
import re
import random
import streamlit as st
from docx import Document
from io import BytesIO

genai.configure(api_key=st.secrets[random.choice(["API_KEY1", "API_KEY2"])])
model = genai.GenerativeModel(model_name="gemini-1.5-flash-002")

if "revision_sheet_created" not in st.session_state:
    st.session_state.revision_sheet_created = False
    st.session_state.revision_sheet = None
    st.session_state.revision_subject = None
    st.session_state.sheet = None
    st.session_state.sheet_already_written = False
    st.toast(f"Bienvenue dans le créateur de fiche de révison {st.session_state.user_id} !")

st.title("📝 Créateur de fiche de révision")

def create_revision_sheet(title, text):
    final_text = re.sub(r'["#*"]', "", text)
    doc = Document()
    doc.add_heading(title, level=1)
    doc.add_paragraph(final_text)
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

def stream_text(text):
    container = st.empty()
    displayed_text = ""
    for char in text:
        displayed_text += char
        st.session_state.sheet = container.markdown(displayed_text)
        time.sleep(0.002)

if not st.session_state.revision_sheet_created:
    st.subheader("Sur quoi veux-tu créer ta fiche de révision ?")
    with st.container(border=True):
        col1, col2 = st.columns(2)
        with col1:
            st.session_state.revision_subject = st.selectbox("📚 **Sélectionne la matière de la fiche de révision :** ", ["Français", "Mathématiques", "Histoire","Géographie","EMC", "Sciences et Vie de la Terre", "Physique Chimie","Technologie", "Anglais","Allemand", "Espagnol"], )
            prompt = st.text_input("📝 **Sujet de la fiche de révision :**", placeholder="Ex : la Seconde Guerre Mondiale")
        with col2:
            difficulty_user = st.text_input("🤯 **Tes difficultés (optionel) :**", placeholder="Ex : les dates")
            st.write("**Prix : 1 ⭐**")
        optional_prompt = ""
        if st.button("📝 Créer la fiche de révsion"):
            if st.session_state.revision_subject and prompt:
                if user_manager.use_credit(user_id=st.session_state.user_id, credits_to_use=1):
                    time.sleep(0.5)
                    if difficulty_user:
                        optional_prompt = f"L'utilisateur a du mal avec : {difficulty_user}. Concentre toi là dessus."
                    with st.spinner("L'EtudIAnt réfléchit..."):
                        response = response = model.generate_content(f"Crée une petite fiche de révision de niveau {user_manager.get_any_user_data(user_id=st.session_state.user_id, column='class_level')} qui peut tenir sur une feuille A4 en {st.session_state.revision_subject} sur le sujet suivant : {prompt}. Sans tableaux. {optional_prompt}")
                    st.session_state.revision_sheet = response.text
                    st.session_state.revision_sheet_created = True
                else:
                    st.error("Pas assez d'Etoiles, revient demain.")
            else:
                st.error("Remplis tous les champs obligatoires.")

if st.session_state.revision_sheet_created:
        with st.container(border=True):
            if not st.session_state.sheet_already_written:
                stream_text(text=st.session_state.revision_sheet)
                st.session_state.sheet_already_written = True
            else:
                st.write(st.session_state.revision_sheet)
            word_folder = create_revision_sheet(title=f"Fiche de révision en {st.session_state.revision_subject}", text=st.session_state.revision_sheet)
            st.download_button(
                label="📥 Télécharger la fiche de révision en format Word",
                data=word_folder,
                file_name=f"{st.session_state.revision_subject}.docx",
                mime="text/docx"
            )
            if st.button("📝 Créer une autre fiche de révision"):
                del st.session_state.revision_sheet_created
                st.rerun()